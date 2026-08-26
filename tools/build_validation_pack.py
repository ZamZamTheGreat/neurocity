from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "deliverables" / "NeuroCity_Validation_and_Outreach_Pack.docx"
ASSETS = ROOT / "deliverables" / "assets"

INK = "101218"
CHARCOAL = "20242B"
GOLD = "D59B32"
VIOLET = "7457FF"
PALE = "F4F2FF"
LIGHT = "F2F4F7"
MID = "667085"
WHITE = "FFFFFF"
RED = "9B1C1C"
GREEN = "176B45"
BLUE = "2463EB"
USABLE_DXA = 9360


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths):
    assert sum(widths) == USABLE_DXA
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(USABLE_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_font(run, size=11, bold=False, color=INK, italic=False, name="Aptos"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, color, before, after in (
        ("Heading 1", 18, INK, 14, 7),
        ("Heading 2", 13, VIOLET, 10, 5),
        ("Heading 3", 11.5, CHARCOAL, 8, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Aptos Display")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos Display")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_style in ("List Bullet", "List Number"):
        style = doc.styles[list_style]
        style.font.name = "Aptos"
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.15


def add_running_furniture(section):
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("NEUROCITY  /  PILOT VALIDATION")
    set_font(r, 8, True, MID)
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run("Confidential working material | 9 August 2026")
    set_font(r, 8, False, MID)


def add_title(doc, title, subtitle=None, kicker=None):
    if kicker:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run(kicker.upper())
        set_font(r, 9, True, GOLD)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(title)
    set_font(r, 28, True, INK, name="Aptos Display")
    if subtitle:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(14)
        r = p.add_run(subtitle)
        set_font(r, 13, False, MID)


def add_callout(doc, label, text, fill=PALE, accent=VIOLET):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [USABLE_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label.upper() + "  ")
    set_font(r, 9, True, accent)
    r = p.add_run(text)
    set_font(r, 10.5, False, INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_bullets(doc, items, numbered=False):
    style = "List Number" if numbered else "List Bullet"
    for item in items:
        p = doc.add_paragraph(style=style)
        p.add_run(item)


def add_key_value_table(doc, rows, label_width=2300, header=None):
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    if header:
        cells = table.add_row().cells
        cells[0].text, cells[1].text = header
        for cell in cells:
            set_cell_shading(cell, CHARCOAL)
            for run in cell.paragraphs[0].runs:
                set_font(run, 9.5, True, WHITE)
        set_repeat_table_header(table.rows[0])
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        set_cell_shading(cells[0], LIGHT)
        for run in cells[0].paragraphs[0].runs:
            set_font(run, 9.5, True, CHARCOAL)
        for run in cells[1].paragraphs[0].runs:
            set_font(run, 9.5, False, INK)
    set_table_geometry(table, [label_width, USABLE_DXA - label_width])
    return table


def add_three_col_table(doc, headers, rows, widths=(2200, 3600, 3560)):
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for idx, heading in enumerate(headers):
        table.rows[0].cells[idx].text = heading
        set_cell_shading(table.rows[0].cells[idx], CHARCOAL)
        for run in table.rows[0].cells[idx].paragraphs[0].runs:
            set_font(run, 9, True, WHITE)
    set_repeat_table_header(table.rows[0])
    for row_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            if row_idx % 2:
                set_cell_shading(cells[idx], "FAFAFB")
            for run in cells[idx].paragraphs[0].runs:
                set_font(run, 9, False, INK)
    set_table_geometry(table, list(widths))
    return table


def page_break(doc):
    doc.add_page_break()


def build():
    doc = Document()
    style_document(doc)
    add_running_furniture(doc.sections[0])

    # Cover
    doc.add_paragraph().paragraph_format.space_after = Pt(36)
    add_title(doc, "NeuroCity", "Pilot validation and outreach pack", "Windhoek digital mall")
    add_callout(doc, "Purpose", "Ready-to-use material for the LightWork Clothing discovery meeting and payment-provider conversations with FNB Namibia, Nedbank Namibia, and PayToday.")
    doc.add_paragraph().paragraph_format.space_after = Pt(14)
    add_key_value_table(doc, [
        ("Pilot city", "Windhoek"),
        ("Pilot categories", "Fashion; Beauty and personal care; Gifts, home and living"),
        ("First merchant prospect", "LightWork Clothing - Baines Centre, Pioneerspark"),
        ("Pilot commercial model", "No charge during testing; monthly merchant subscription after pilot; transaction commission deferred"),
        ("Customer payment", "Online payment and pay on collection"),
        ("Fulfillment", "Merchant-managed delivery and pickup initially"),
        ("Platform relationship", "Independent from NeuroEstates"),
    ])
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    r = p.add_run("Prepared for Sergej Witbooi")
    set_font(r, 10, True, MID)

    # Decision record
    page_break(doc)
    add_title(doc, "Pilot decision record", "The agreed boundaries that all research and wireframes should follow.", "01 / decisions")
    add_three_col_table(doc, ("Decision", "Confirmed direction", "Implication"), [
        ("Name", "NeuroCity", "Develop an independent public identity and domain strategy."),
        ("Geography", "Windhoek", "Use zone-based local delivery and store pickup for V1."),
        ("Categories", "Fashion; Beauty/personal care; Gifts/home/living", "Build category attributes without grocery or regulated-goods complexity."),
        ("Testing price", "No merchant charge", "Use a short, defined pilot with explicit exit criteria."),
        ("Post-pilot revenue", "Monthly subscription", "Validate willingness to pay and tier requirements during interviews."),
        ("Commission", "Not included for now", "Keep commission fields in architecture but do not promise a rate."),
        ("Payments", "Online plus pay on collection", "Require provider validation and merchant-level eligibility rules."),
        ("Delivery", "Merchant managed initially", "Capture zones, fees, SLA, proof, and contact process per merchant."),
        ("NeuroEstates", "Separate product", "No shared identity or infrastructure dependency in the pilot."),
    ])
    add_callout(doc, "Working assumption", "A mall-wide basket can group products visually, but each merchant order is paid independently until a provider confirms compliant marketplace settlement.")

    # Merchant introduction
    page_break(doc)
    add_title(doc, "Merchant introduction", "A short explanation to use before the LightWork discovery conversation.", "02 / merchant pack")
    doc.add_heading("Suggested opening", level=1)
    p = doc.add_paragraph()
    p.add_run("NeuroCity is a Windhoek-first digital mall where local stores keep their own identity while customers discover products, ask intelligent shopping questions, and order for pickup or delivery in one place. We are designing a controlled pilot with a small group of merchants. There is no charge during testing. The objective is to learn what genuinely helps merchants sell online before final subscription plans are introduced.")
    doc.add_heading("What LightWork would receive in the pilot", level=2)
    add_bullets(doc, [
        "A branded digital storefront inside NeuroCity.",
        "Structured product, variant, price, and availability presentation.",
        "Customer checkout with online payment and approved pay-on-collection options.",
        "Pickup at Baines Centre and merchant-managed delivery configuration.",
        "A merchant dashboard for products, stock, and orders.",
        "A catalogue-grounded shopping assistant that only recommends published LightWork products.",
        "Pilot analytics and direct feedback on customer demand.",
    ])
    doc.add_heading("What we need from LightWork", level=2)
    add_bullets(doc, [
        "One responsible person for catalogue and order decisions.",
        "Accurate products, variants, prices, stock status, and product images.",
        "Pickup hours, delivery coverage, fees, preparation times, and returns policy.",
        "Fast feedback during testing and agreement to fulfill accepted pilot orders.",
    ])
    add_callout(doc, "Do not promise yet", "A launch date, guaranteed sales, a commission rate, one-payment multi-store settlement, or NeuroCity-managed delivery.", fill="FFF7E8", accent=GOLD)

    # LightWork profile
    page_break(doc)
    add_title(doc, "LightWork Clothing discovery brief", "Known facts, evidence, and items that require confirmation.", "03 / first pilot prospect")
    add_key_value_table(doc, [
        ("Business", "LightWork Clothing"),
        ("Category", "Fashion / streetwear"),
        ("Contact", "Zephan Stadhauer"),
        ("Phone", "081 495 3446"),
        ("Email", "lightworkclothing.na@gmail.com"),
        ("Website", "https://lightworkclothing.com"),
        ("Relationship", "Personal connection through the same gym"),
        ("Branches", "One"),
        ("Pickup", "Baines Centre, Pioneerspark, Windhoek"),
        ("Current online sales", "Website exists; Zephan indicated it is not a proper transactional store"),
    ])
    doc.add_heading("Evidence supplied", level=2)
    add_three_col_table(doc, ("Collection/product", "Visible evidence", "Confirm before publishing"), [
        ("Crown V1 Cuffed Tracksuit", "Black tracksuit; image shows N$1,249.99 and historic preorder language.", "Current price, sale mode, component/size variants, stock, latest images."),
        ("Metallic 23 Longsleeve - Maroon", "Maroon long-sleeve; image states 200 GSM and 100% cotton.", "Current name, material claim, price, sizes, stock, wash/care details."),
        ("Majesteric Edition", "Black zip hoodies with blue, red, pink, and purple butterfly artwork.", "Exact SKUs, current availability, prices, sizes, edition/drop status."),
        ("Esoteric", "T-shirt/short sets shown in several colours.", "Whether sold as set or separates, official colour names, prices, sizes, stock."),
    ])
    add_callout(doc, "Data rule", "Promotional artwork is not treated as current inventory. Every customer-facing claim requires merchant confirmation.")

    # Interview guide
    page_break(doc)
    add_title(doc, "LightWork discovery interview", "A 45-minute conversation guide. Record answers in the companion workbook or meeting notes.", "04 / interview")
    add_three_col_table(doc, ("Time", "Topic", "Outcome"), [
        ("0-5 min", "Context and pilot", "Shared understanding; permission to continue."),
        ("5-15 min", "Products and catalogue", "SKU/variant structure and available assets."),
        ("15-25 min", "Stock and orders", "Current process, ownership, response times."),
        ("25-33 min", "Pickup, delivery, returns", "Fulfillment rules and customer promises."),
        ("33-40 min", "Storefront and AI", "Brand controls and common customer questions."),
        ("40-45 min", "Pilot fit and next steps", "Commitment, owner, dates, missing items."),
    ], widths=(1200, 3200, 4960))
    doc.add_heading("Questions", level=1)
    questions = [
        "Which products are actively sold today, and which are past or future drops?",
        "How do you define products versus variants - for example colour, size, set, or separate piece?",
        "Where do you currently record prices and stock? How often is stock accurate?",
        "Who can update products and respond when an order arrives?",
        "How quickly can you accept or reject an order during opening hours?",
        "Which sizes and colours should customers be allowed to select?",
        "Can stock be reserved briefly during online checkout?",
        "What are the exact pickup days, hours, instructions, and preparation time at Baines Centre?",
        "Which Windhoek areas do you deliver to, what do you charge, and who performs delivery?",
        "When is pay on collection acceptable? Should items be held before payment?",
        "What cancellations, exchanges, and returns do you currently allow?",
        "Which questions do customers ask repeatedly before buying?",
        "What parts of the existing LightWork identity must never be changed?",
        "Would you use a monthly subscription after testing? What would it need to include?",
        "What outcome would make the pilot worthwhile after 30-60 days?",
    ]
    add_bullets(doc, questions, numbered=True)
    doc.add_heading("Close with", level=2)
    add_bullets(doc, [
        "Confirm the pilot decision maker and day-to-day operator.",
        "Agree on a date for receiving the first completed catalogue workbook.",
        "Agree on which 5-10 products will form the initial demo catalogue.",
        "Request written permission to use brand and product media in the private pilot prototype.",
        "Schedule a 30-minute storefront review after the first wireframe is ready.",
    ])

    # Catalogue checklist
    page_break(doc)
    add_title(doc, "Merchant onboarding checklist", "Material required before LightWork can enter a controlled live pilot.", "05 / onboarding")
    add_three_col_table(doc, ("Area", "Required", "Status for LightWork"), [
        ("Business", "Registered/trading details, contact, branch, operating hours", "Partially known"),
        ("Brand", "Approved logo, colours, store description, social links", "Logo supplied; remaining items needed"),
        ("Catalogue", "5-10 active products with variants, price, SKU, description", "Draft evidence only"),
        ("Media", "Clean product images and permission to use them", "Promotional images supplied; approval needed"),
        ("Stock", "On-hand or availability mode, update owner and frequency", "Needed"),
        ("Pickup", "Exact location, hours, lead time, verification process", "Location known; details needed"),
        ("Delivery", "Zones, fee, timing, contact and proof method", "Needed"),
        ("Policies", "Cancellation, exchange, returns, privacy/contact consent", "Needed"),
        ("Orders", "Responsible staff, notification channel, response SLA", "Needed"),
        ("Pilot", "Participation approval, success measure, review date", "Needed"),
    ])
    doc.add_heading("Product photo standard", level=2)
    add_bullets(doc, [
        "At least one clear front image; back and detail views strongly preferred.",
        "Consistent crop, neutral or intentional background, and no unreadable small text.",
        "One product/variant represented accurately; avoid using a collage as the only image.",
        "Minimum recommended 1200 x 1200 pixels for square product cards.",
        "No outdated price/date text embedded in the primary product image.",
        "Merchant confirms ownership or permission to use each asset.",
    ])

    # Pilot expectations
    page_break(doc)
    add_title(doc, "Pilot expectations to discuss", "Commercial points for a later participation agreement; not a legal contract.", "06 / pilot terms")
    add_three_col_table(doc, ("Topic", "Proposed pilot position", "Confirm with merchant"), [
        ("Fees", "No subscription or transaction commission during controlled testing.", "Pilot start/end and any payment-processing fees."),
        ("Catalogue", "Merchant is responsible for accuracy and authorizes approved display.", "Update frequency and content owner."),
        ("Orders", "Merchant accepts only fulfillable orders and responds within an agreed SLA.", "Hours and target response time."),
        ("Stock", "Online stock or confirmation status must reflect operational reality.", "Exact stock versus confirm-first mode."),
        ("Payment", "Provider processes online payments; pay on collection is merchant-configured.", "Refund and no-show process."),
        ("Fulfillment", "Merchant controls delivery; NeuroCity displays agreed information and status.", "Zones, fees, liability and proof."),
        ("Returns", "Published policy applies subject to the mall's minimum framework and law.", "Exchange/return conditions."),
        ("Brand", "Merchant retains brand; NeuroCity may use approved assets for the pilot.", "Scope and revocation."),
        ("Data", "Order data is used to provide and improve the pilot; marketing is separately consented.", "Access and retention notice."),
        ("Exit", "Either party can leave the test under a defined wind-down process.", "Notice and outstanding-order handling."),
    ])
    add_callout(doc, "Before live money", "Have the final merchant agreement, customer terms, privacy notice, returns framework, tax treatment, and payment flow reviewed by qualified Namibian advisers.", fill="FFF1F1", accent=RED)

    # Provider email
    page_break(doc)
    add_title(doc, "Payment-provider enquiry", "Copy-ready email for FNB Namibia, Nedbank Namibia, or PayToday.", "07 / payments")
    add_key_value_table(doc, [
        ("Subject", "NeuroCity pilot - e-commerce and future marketplace payment capability"),
        ("To", "[Relationship manager / merchant services contact]"),
        ("From", "Sergej Witbooi / NeuroCity"),
    ])
    p = doc.add_paragraph()
    p.add_run("Dear [Name],\n\n")
    p.add_run("We are preparing a controlled Windhoek pilot for NeuroCity, a digital mall that gives participating local merchants branded storefronts, product catalogues, ordering, pickup, and merchant-managed delivery. The pilot will begin with a small set of merchants across fashion, beauty/personal care, and gifts/home and living.\n\n")
    p.add_run("We would like to understand the most appropriate compliant payment structure for the pilot. Initially, customers should be able to pay online or pay on collection where the merchant allows it. Each merchant portion can be treated as a separate order and payment. Longer term, we want to assess whether your platform supports one customer payment with marketplace split settlement, merchant onboarding, platform fees, partial refunds, chargebacks, and reconciliation.\n\n")
    p.add_run("Please arrange a discussion with the relevant e-commerce or merchant-services specialist and provide written responses to the attached capability checklist. We already maintain business banking relationships in Namibia and would like to evaluate technical capability, compliance responsibilities, pricing, onboarding time, and settlement reporting.\n\n")
    p.add_run("No card credentials or sensitive customer payment data will be stored by NeuroCity; we expect to use provider-hosted payment collection and verified transaction callbacks.\n\n")
    p.add_run("Kind regards,\nSergej Witbooi\nNeuroCity\n[preferred phone] | [preferred email]")

    # Provider questions
    page_break(doc)
    add_title(doc, "Payment capability checklist", "Ask for written answers and a technical follow-up.", "08 / provider comparison")
    add_three_col_table(doc, ("Area", "Question", "Evidence requested"), [
        ("Regulatory role", "Which licensed entity provides acquiring/payment services for this flow?", "Licence/authorization and contracting entity."),
        ("Pilot structure", "Can separate merchant payments be presented in one guided checkout?", "Recommended architecture and constraints."),
        ("Marketplace", "Can one payment be split among merchants and NeuroCity later?", "Product name, agreement model, flow diagram."),
        ("Onboarding", "Who performs merchant KYC and how long does activation take?", "Required documents and SLA."),
        ("Methods", "Which cards, EFT/mobile methods, and 3DS flows support NAD?", "Current supported-method list."),
        ("Integration", "Are hosted fields/redirect, API, sandbox, webhooks, and idempotency available?", "Developer documentation and sandbox access."),
        ("Refunds", "Can we initiate full and partial refunds by API?", "Limits, approval flow, timing, fees."),
        ("Disputes", "Who manages chargebacks and bears financial responsibility?", "Process, evidence window, fee schedule."),
        ("Settlement", "How are merchants paid and how are platform fees represented?", "Settlement cadence and sample report."),
        ("Reconciliation", "Can every charge/refund/settlement be exported with stable IDs?", "Sample CSV/API schema."),
        ("Security", "Which PCI scope remains with NeuroCity? How are webhooks signed?", "AOC/attestation guidance and security docs."),
        ("Commercials", "What setup, monthly, transaction, refund, dispute, and reserve costs apply?", "Formal pricing proposal."),
        ("Support", "What incident, escalation, and production support is provided?", "Support hours, channels, and SLA."),
    ], widths=(1600, 4500, 3260))

    # comparison scorecard
    page_break(doc)
    add_title(doc, "Provider evaluation scorecard", "Use the workbook or this compact view after meetings.", "09 / decision method")
    add_three_col_table(doc, ("Criterion", "Weight", "Pass condition"), [
        ("Regulatory and contractual fit", "Mandatory", "Provider confirms the proposed flow in writing."),
        ("Separate merchant payment support", "15%", "Reliable merchant-scoped charges and refunds."),
        ("Future marketplace settlement", "15%", "Documented split settlement or viable roadmap."),
        ("Integration and sandbox", "15%", "Hosted checkout, APIs, signed webhooks, test environment."),
        ("Refunds and disputes", "10%", "Partial refunds, auditable chargeback process."),
        ("Merchant onboarding", "10%", "Practical KYC and activation for small merchants."),
        ("Reconciliation", "10%", "Stable transaction IDs and useful statements/exports."),
        ("Payment-method fit", "10%", "NAD and relevant local customer methods."),
        ("Commercial terms", "10%", "Affordable pilot and sustainable scaled pricing."),
        ("Support and reliability", "5%", "Named escalation path and credible service levels."),
    ], widths=(3500, 1300, 4560))
    add_callout(doc, "Decision rule", "Do not select solely on transaction price. Regulatory fit, refund handling, reconciliation, and operational support are launch-critical.")

    # visual identity
    page_break(doc)
    add_title(doc, "NeuroCity visual direction", "Premium modern mall + friendly local marketplace + restrained AI energy.", "10 / identity proposal")
    add_key_value_table(doc, [
        ("Brand idea", "A connected city of trusted local storefronts"),
        ("Personality", "Confident, useful, welcoming, intelligent, locally grounded"),
        ("Primary ink", "Midnight Ink #101218 - premium foundation"),
        ("Urban charcoal", "#20242B - navigation and structural surfaces"),
        ("Namib gold", "#D59B32 - warmth, local character, featured moments"),
        ("Neural violet", "#7457FF - AI actions and interactive emphasis"),
        ("Cloud", "#F2F4F7 - open, clean commerce canvas"),
        ("Typography", "Aptos/Inter-like geometric sans for product; expressive merchant typography remains inside storefronts"),
    ])
    doc.add_heading("Design principles", level=2)
    add_bullets(doc, [
        "NeuroCity provides the reliable mall shell; merchants retain their own voice and product imagery.",
        "AI is visible through precise violet accents and helpful language, not sci-fi decoration.",
        "Gold is selective - featured stores, trusted highlights, and city warmth - never used everywhere.",
        "Mobile layouts prioritize product, price, availability, and fulfillment over visual spectacle.",
        "Photography should feel real and local, with diverse Windhoek contexts where appropriate.",
    ])
    doc.add_heading("Initial tagline routes", level=2)
    add_bullets(doc, [
        "Windhoek shops here.",
        "Your city. Your stores. One place.",
        "Discover local. Shop intelligently.",
    ])
    add_callout(doc, "Recommended working line", "Your city. Your stores. One place.", fill="FFF7E8", accent=GOLD)

    # image references
    page_break(doc)
    add_title(doc, "LightWork visual reference", "Supplied materials for a private pilot storefront concept; final use requires merchant approval.", "11 / merchant identity")
    logo = ASSETS / "lightwork-logo.png"
    if logo.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(logo), width=Inches(1.2))
    img = ASSETS / "lightwork-crown-v1.png"
    if img.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(img), width=Inches(3.5))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Reference: Crown V1 promotional artwork supplied by Sergej. Historic date/preorder text must not be used as current sales information.")
    set_font(r, 8.5, False, MID, italic=True)
    add_callout(doc, "Storefront treatment", "Use LightWork's black-and-white identity inside its store, framed by NeuroCity's neutral navigation. Product cards should use clean individual images rather than promotional collages where possible.")

    # next steps
    page_break(doc)
    add_title(doc, "Action plan", "The shortest route from validation to wireframes and a build decision.", "12 / next steps")
    add_three_col_table(doc, ("Owner", "Action", "Output"), [
        ("Sergej", "Send the merchant introduction and schedule the 45-minute LightWork conversation.", "Meeting date and permission to explore pilot."),
        ("Sergej + Zephan", "Complete the interview and choose 5-10 active products.", "Confirmed operating answers and pilot catalogue."),
        ("Zephan", "Complete missing catalogue, stock, fulfillment, and policy fields.", "Usable LightWork onboarding workbook."),
        ("Sergej", "Send the provider email/checklist through FNB, Nedbank, and PayToday contacts.", "Written capability and pricing responses."),
        ("NeuroCity", "Compare providers using mandatory gates and weighted criteria.", "Payment flow decision."),
        ("NeuroCity", "Create customer, LightWork, merchant-dashboard, and admin wireframes.", "Testable low-fidelity prototype."),
        ("Sergej + pilot merchant", "Review the prototype against real order/fulfillment scenarios.", "Signed-off pilot workflow and revised backlog."),
    ], widths=(1700, 4600, 3060))
    doc.add_heading("Definition of validation complete", level=2)
    add_bullets(doc, [
        "LightWork confirms interest, a responsible operator, and a starter catalogue.",
        "Pickup, delivery, order response, returns, and pay-on-collection rules are documented.",
        "At least one payment provider confirms a compliant pilot flow and integration path in writing.",
        "The initial subscription hypothesis and pilot success measure are recorded.",
        "Wireframes can be tested using real products and realistic operational states.",
    ])
    add_callout(doc, "Current status", "Materials prepared. Next external dependency: schedule the LightWork meeting and send payment-provider enquiries.", fill="EAF7F0", accent=GREEN)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
